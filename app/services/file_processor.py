"""
File Processor Service - Extract text from various file formats
"""

import logging
import os
from typing import Dict, Optional
import aiofiles
from PyPDF2 import PdfReader
from docx import Document
import markdown
from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)


class FileProcessor:
    """Service for processing and extracting text from uploaded files"""
    
    @staticmethod
    async def process_pdf(file_path: str) -> Dict:
        """
        Extract text from PDF file
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Dict with text_content, word_count, and metadata
        """
        try:
            reader = PdfReader(file_path)
            text_content = ""
            metadata = {
                "pages": len(reader.pages),
                "title": reader.metadata.title if reader.metadata and reader.metadata.title else None,
                "author": reader.metadata.author if reader.metadata and reader.metadata.author else None,
            }
            
            # Extract text from all pages
            for page in reader.pages:
                text_content += page.extract_text() + "\n"
            
            word_count = len(text_content.split())
            
            return {
                "text_content": text_content.strip(),
                "word_count": word_count,
                "metadata": metadata,
                "mime_type": "application/pdf"
            }
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            raise ValueError(f"Failed to process PDF: {str(e)}")
    
    @staticmethod
    async def process_docx(file_path: str) -> Dict:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dict with text_content, word_count, and metadata
        """
        try:
            doc = Document(file_path)
            text_content = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
                text_content += "\n"
            
            word_count = len(text_content.split())
            
            metadata = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
            }
            
            # Try to get core properties
            try:
                if doc.core_properties.title:
                    metadata["title"] = doc.core_properties.title
                if doc.core_properties.author:
                    metadata["author"] = doc.core_properties.author
            except:
                pass
            
            return {
                "text_content": text_content.strip(),
                "word_count": word_count,
                "metadata": metadata,
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            }
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            raise ValueError(f"Failed to process DOCX: {str(e)}")
    
    @staticmethod
    async def process_text(file_path: str) -> Dict:
        """
        Extract text from plain text file
        
        Args:
            file_path: Path to text file
            
        Returns:
            Dict with text_content, word_count, and metadata
        """
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text_content = await f.read()
            
            word_count = len(text_content.split())
            
            return {
                "text_content": text_content.strip(),
                "word_count": word_count,
                "metadata": {},
                "mime_type": "text/plain"
            }
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {e}")
            raise ValueError(f"Failed to process text file: {str(e)}")
    
    @staticmethod
    async def process_markdown(file_path: str) -> Dict:
        """
        Extract text from Markdown file
        
        Args:
            file_path: Path to markdown file
            
        Returns:
            Dict with text_content, word_count, and metadata
        """
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                md_content = await f.read()
            
            # Convert markdown to HTML then extract text
            html_content = markdown.markdown(md_content)
            soup = BeautifulSoup(html_content, 'html.parser')
            text_content = soup.get_text()
            
            word_count = len(text_content.split())
            
            return {
                "text_content": text_content.strip(),
                "word_count": word_count,
                "metadata": {"format": "markdown"},
                "mime_type": "text/markdown"
            }
        except Exception as e:
            logger.error(f"Error processing markdown file {file_path}: {e}")
            raise ValueError(f"Failed to process markdown file: {str(e)}")
    
    @staticmethod
    async def process_html(file_path: str) -> Dict:
        """
        Extract text from HTML file
        
        Args:
            file_path: Path to HTML file
            
        Returns:
            Dict with text_content, word_count, and metadata
        """
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                html_content = await f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Extract title
            title = soup.title.string if soup.title else None
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            text_content = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text_content.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text_content = '\n'.join(chunk for chunk in chunks if chunk)
            
            word_count = len(text_content.split())
            
            return {
                "text_content": text_content.strip(),
                "word_count": word_count,
                "metadata": {"title": title} if title else {},
                "mime_type": "text/html"
            }
        except Exception as e:
            logger.error(f"Error processing HTML file {file_path}: {e}")
            raise ValueError(f"Failed to process HTML file: {str(e)}")
    
    @staticmethod
    def detect_mime_type(filename: str) -> str:
        """
        Detect MIME type from filename extension
        
        Args:
            filename: Name of the file
            
        Returns:
            MIME type string
        """
        extension = os.path.splitext(filename)[1].lower()
        
        mime_types = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.txt': 'text/plain',
            '.md': 'text/markdown',
            '.html': 'text/html',
            '.htm': 'text/html',
        }
        
        return mime_types.get(extension, 'application/octet-stream')
    
    @staticmethod
    async def extract_text(file_path: str, mime_type: str) -> Dict:
        """
        Extract text from file based on MIME type
        
        Args:
            file_path: Path to file
            mime_type: MIME type of the file
            
        Returns:
            Dict with text_content, word_count, and metadata
        """
        try:
            if mime_type == 'application/pdf':
                return await FileProcessor.process_pdf(file_path)
            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                return await FileProcessor.process_docx(file_path)
            elif mime_type == 'text/plain':
                return await FileProcessor.process_text(file_path)
            elif mime_type == 'text/markdown':
                return await FileProcessor.process_markdown(file_path)
            elif mime_type in ['text/html', 'application/xhtml+xml']:
                return await FileProcessor.process_html(file_path)
            else:
                raise ValueError(f"Unsupported file type: {mime_type}")
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            raise


# Singleton instance
file_processor = FileProcessor()
